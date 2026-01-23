"""Document parsing service using Docling for advanced PDF parsing."""

import base64
import binascii
import os
import tempfile

import structlog

logger = structlog.get_logger()

# Constants
BASE64_THRESHOLD = 0.1  # Consider base64 if <10% non-base64 chars


class DocumentParser:
    """Service for parsing documents using Docling."""

    def __init__(self) -> None:
        """Initialize document parser."""
        self.docling_available = False
        self._init_docling()

    def _init_docling(self) -> None:
        """Initialize Docling if available."""
        try:
            from docling.document_converter import DocumentConverter

            self.converter = DocumentConverter()
            self.docling_available = True
            logger.info("docling_initialized", available=True)
        except ImportError:
            logger.warning(
                "docling_not_available",
                message="Docling is not installed. Use: pip install docling",
            )
            self.docling_available = False
        except Exception as e:
            logger.warning("docling_init_failed", error=str(e))
            self.docling_available = False

    def _looks_like_base64(self, s: str, threshold: float = BASE64_THRESHOLD) -> bool:
        """
        Check if string looks like base64 encoded data.

        Args:
            s: String to check
            threshold: Threshold for non-base64 characters (default BASE64_THRESHOLD = 10%)

        Returns:
            True if string appears to be base64 encoded
        """
        if len(s) < 100:
            return False
        base64_chars = set("ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789+/=")
        non_base64 = sum(1 for c in s if c not in base64_chars and not c.isspace())
        return non_base64 < len(s) * threshold

    def parse_document(
        self,
        content: str,
        filename: str,
        content_type: str,
    ) -> str:
        """
        Parse document content using Docling for PDFs, python-docx for DOCX, fallback for others.

        Args:
            content: Document content (can be base64 encoded or plain text)
            filename: Original filename
            content_type: MIME type of the document

        Returns:
            Extracted text content
        """
        # Check if it's a PDF that should use Docling
        is_pdf = (
            filename.lower().endswith(".pdf")
            or content_type == "application/pdf"
            or "pdf" in content_type.lower()
        )

        # Check if it's a DOCX file
        is_docx = (
            filename.lower().endswith(".docx")
            or content_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or "docx" in content_type.lower()
            or "word" in content_type.lower()
        )

        if is_pdf and self.docling_available:
            return self._parse_with_docling(content, filename)
        elif is_docx:
            return self._parse_docx(content, filename)
        else:
            # Fallback to simple text extraction
            return self._parse_simple(content, filename, content_type)

    def _parse_with_docling(self, content: str, filename: str) -> str:
        """
        Parse PDF using Docling.

        Args:
            content: Document content (base64 or raw bytes)
            filename: Original filename

        Returns:
            Extracted text content
        """
        try:

            logger.info("parsing_with_docling", filename=filename)

            # Try to decode base64 if needed
            try:
                # Check if content is base64 encoded
                if isinstance(content, str) and not content.startswith("%PDF"):
                    # Try to decode base64
                    try:
                        pdf_bytes = base64.b64decode(content)
                    except Exception:
                        # If not base64, assume it's already raw bytes as string
                        pdf_bytes = (
                            content.encode("latin-1") if isinstance(content, str) else content
                        )
                else:
                    pdf_bytes = content.encode("latin-1") if isinstance(content, str) else content

                # Docling needs a temporary file (doesn't accept BytesIO)
                # Create temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
                    tmp_file.write(pdf_bytes)
                    tmp_path = tmp_file.name

                try:
                    # Convert using Docling with file path
                    result = self.converter.convert(tmp_path)
                finally:
                    # Clean up temporary file
                    try:
                        os.unlink(tmp_path)
                    except Exception:
                        pass

                # Extract text from DoclingDocument
                # Docling returns a DoclingDocument with export_to_markdown() method
                if hasattr(result, "export_to_markdown"):
                    text = result.export_to_markdown()
                elif hasattr(result, "document") and hasattr(result.document, "export_to_markdown"):
                    # Try document.export_to_markdown()
                    text = result.document.export_to_markdown()
                elif hasattr(result, "document") and hasattr(result.document, "export_to_text"):
                    # Try document.export_to_text() if available
                    text = result.document.export_to_text()
                elif hasattr(result, "document"):
                    # Last resort: try to extract text from document structure
                    # Avoid str() as it shows object representation
                    try:
                        # Try to get text content from document items
                        if hasattr(result.document, "items"):
                            text_parts = []
                            for item in result.document.items:
                                if hasattr(item, "text"):
                                    text_parts.append(str(item.text))
                                elif hasattr(item, "content"):
                                    text_parts.append(str(item.content))
                            text = "\n".join(text_parts) if text_parts else ""
                        else:
                            # If no items, try markdown export on document
                            text = (
                                result.document.export_to_markdown()
                                if hasattr(result.document, "export_to_markdown")
                                else ""
                            )
                    except Exception as e:
                        logger.warning("docling_text_extraction_failed", error=str(e))
                        text = ""
                else:
                    # Last fallback - try to get any text representation
                    text = ""

                # Clean up text - remove any remaining object representations
                if text and ("TextItem" in text or "RefItem" in text or "<" in text[:100]):
                    # Text might contain object representations, try to extract actual text
                    import re

                    # Try to extract text from markdown-like content
                    text = re.sub(r"TextItem\([^)]*text=\'([^\']+)\'[^)]*\)", r"\1", text)
                    text = re.sub(r"RefItem\([^)]*\)", "", text)
                    # Remove other object representations
                    text = re.sub(r"<[^>]+>", "", text)

                if not text or len(text.strip()) == 0:
                    logger.warning("docling_no_text_extracted", filename=filename)
                    # Fallback to simple parsing
                    return self._parse_simple(content, filename, "application/pdf")

                logger.info("docling_parse_success", filename=filename, text_length=len(text))
                return text

            except (ValueError, OSError) as e:
                logger.error("docling_parse_error", error=str(e), error_type=type(e).__name__)
                # Fallback to simple parsing
                return self._parse_simple(content, filename, "application/pdf")
            except Exception as e:
                logger.error("docling_unexpected_error", error=str(e), error_type=type(e).__name__)
                # Re-raise unexpected errors
                raise

        except Exception as e:
            logger.error("docling_unavailable", error=str(e), error_type=type(e).__name__)
            return self._parse_simple(content, filename, "application/pdf")

    def _parse_docx(self, content: str, filename: str) -> str:
        """
        Parse DOCX file using python-docx.

        Args:
            content: Document content (base64 encoded)
            filename: Original filename

        Returns:
            Extracted text content
        """
        try:
            import io

            from docx import Document

            logger.info("parsing_docx", filename=filename)

            # Decode base64 if needed
            # DOCX files start with PK (ZIP signature) when decoded
            # Check if content is base64 encoded
            try:
                if isinstance(content, str) and self._looks_like_base64(content):
                    # Try to decode base64
                    try:
                        docx_bytes = base64.b64decode(content, validate=True)
                        # Verify it's actually a DOCX (starts with PK)
                        if not docx_bytes.startswith(b"PK"):
                            logger.warning(
                                "decoded_not_docx", message="Decoded content doesn't look like DOCX"
                            )
                            # Fallback to simple parsing
                            return self._parse_simple(
                                content,
                                filename,
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
                    except (ValueError, binascii.Error) as e:
                        logger.warning(
                            "docx_base64_decode_failed", error=str(e), error_type=type(e).__name__
                        )
                        # Fallback to simple parsing
                        return self._parse_simple(
                            content,
                            filename,
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                elif isinstance(content, str):
                    # Already decoded or plain text - try to encode as bytes
                    docx_bytes = content.encode("latin-1")
                else:
                    docx_bytes = content
            except (OSError, ValueError, UnicodeEncodeError) as e:
                logger.error("docx_decode_error", error=str(e), error_type=type(e).__name__)
                # Fallback to simple parsing
                return self._parse_simple(
                    content,
                    filename,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

            # Parse DOCX
            doc = Document(io.BytesIO(docx_bytes))

            # Extract text from all paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            text = "\n\n".join(text_parts)

            if not text or len(text.strip()) == 0:
                logger.warning("docx_no_text_extracted", filename=filename)
                return ""

            logger.info("docx_parse_success", filename=filename, text_length=len(text))
            return text

        except ImportError:
            logger.warning(
                "python_docx_not_available",
                message="python-docx not installed. Install with: pip install python-docx",
            )
            # Fallback to simple extraction
            return self._parse_simple(
                content,
                filename,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except (ValueError, OSError) as e:
            logger.error("docx_parse_error", error=str(e), error_type=type(e).__name__)
            # Fallback to simple extraction
            return self._parse_simple(
                content,
                filename,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception as e:
            logger.error("docx_unexpected_error", error=str(e), error_type=type(e).__name__)
            # Re-raise unexpected errors
            raise

    def _parse_simple(self, content: str, filename: str, content_type: str) -> str:
        """
        Simple text extraction fallback.

        Args:
            content: Document content
            filename: Original filename
            content_type: MIME type

        Returns:
            Extracted text content
        """
        logger.info("using_simple_parser", filename=filename, content_type=content_type)

        # If it's already text, return as is
        if isinstance(content, str):
            # Only try base64 decode if content looks like base64
            # Base64 strings are typically longer and contain only base64 characters
            if self._looks_like_base64(content):
                try:
                    decoded = base64.b64decode(content, validate=True)
                    # Check if decoded content is PDF
                    if decoded.startswith(b"%PDF"):
                        # PDF binary - can't extract text simply
                        logger.warning(
                            "pdf_binary_detected",
                            message="Binary PDF detected. Install Docling to extract text.",
                        )
                        return ""
                    else:
                        # Try to decode as UTF-8
                        return decoded.decode("utf-8", errors="ignore")
                except Exception:
                    # Base64 decode failed, return as plain text
                    logger.warning(
                        "base64_decode_failed",
                        message="Failed to decode base64, treating as plain text",
                    )
                    return content
            else:
                # Doesn't look like base64, return as plain text
                return content

        return str(content) if content else ""
